import io
import os
import re
import uuid
from typing import Any

import pymorphy3
import pytesseract
from PIL import Image, UnidentifiedImageError
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import _Cell
from docx.text.run import Run
from dotenv import load_dotenv
from nltk.corpus import wordnet
from pymorphy3.analyzer import Parse
from pymorphy3.units.by_analogy import KnownSuffixAnalyzer

from config import print

load_dotenv()
TESSERACT_PATH = os.getenv("TESSERACT_PATH")
pytesseract.pytesseract.tesseract_cmd = TESSERACT_PATH

ALLOWED_TEXT_EXT = {'docx', 'doc', 'rtf', 'pdf', 'txt'}
ALLOWED_SPREADSHEET_EXT = {'xls', 'xlsx'}
ALLOWED_EXTENSIONS = ALLOWED_TEXT_EXT.union(ALLOWED_SPREADSHEET_EXT)

from config import spreadsheet_upload_folder as spreadsheet_uf
from config import text_upload_folder as text_uf

morph = pymorphy3.MorphAnalyzer()


def allowed_file(filename):
    return '.' in filename and get_file_format(filename) in ALLOWED_EXTENSIONS


def generate_new_filename(filename: str):
    return str(uuid.uuid4()) + '.' + get_file_format(filename)


def is_filename_secured(filename):
    if not allowed_file(filename):
        return False
    if any(c in filename for c in ['/', '\\', '.py', '../']):
        return False
    return True


def save_uploaded_file(file):
    if get_file_format(file.filename) in ALLOWED_TEXT_EXT:
        path = os.path.join(text_uf, file.filename)
    else:
        path = os.path.join(spreadsheet_uf, file.filename)
    file.save(path)
    print(f"Upload saved in: {path}")
    return file.filename


def get_uploaded_file(filename):
    if get_file_format(filename) in ALLOWED_TEXT_EXT:
        path = os.path.join(text_uf, filename)
    elif get_file_format(filename) in ALLOWED_SPREADSHEET_EXT:
        path = os.path.join(spreadsheet_uf, filename)
    else:
        return None

    if not os.path.exists(path):
        return None

    return path


def get_file_format(filename):
    try:
        return filename.rsplit('.', 1)[1].lower()
    except IndexError:
        return ''


def create_document_and_map(text):
    doc = Document()
    index_map = dict()
    paragraph = doc.add_paragraph()
    for char_index, char in enumerate(text):
        if char == '\n':
            paragraph = doc.add_paragraph()
        else:
            try:
                run = paragraph.add_run(char)
            except ValueError:
                run = paragraph.add_run(' ')
            index_map[char_index] = run

    return doc, index_map


def get_image_text_rotate(image_bytes):
    try:
        img = Image.open(io.BytesIO(image_bytes))
    except UnidentifiedImageError:
        print(f"Something went wrong while reading an image (UnidentifiedImageError). Skipped.")
        return ""
    best_text_probability = 0
    best_image_part = ''
    for lang in ['rus'] * 4 + ['eng'] * 4:
        print(f"lang={lang}")
        image_part = pytesseract.image_to_string(img, lang=lang)
        # если текст плохой, то поворачиваем картинку и распознаём заново
        text_probability = get_text_probability(image_part, lang)
        print(f"text_probability={text_probability}")
        if text_probability is None:
            break
        if text_probability > 0.7:
            best_image_part = image_part
            break
        elif best_text_probability < text_probability:
            best_text_probability = text_probability
            best_image_part = image_part

        print("rotating...")
        img = img.rotate(90, expand=True)
    return best_image_part


def get_image_text(image_bytes):
    try:
        img = Image.open(io.BytesIO(image_bytes))
        # img.show()
    except UnidentifiedImageError:
        print(f"Something went wrong while reading an image (UnidentifiedImageError). Skipped.")
        return ""
    for lang in ['rus'] + ['eng']:
        print(f"lang={lang}")
        image_part = pytesseract.image_to_string(img, lang=lang)
        text_probability = get_text_probability(image_part, lang)
        print(f"text_probability={text_probability}")
    return text_probability


def get_text_probability(text, lang='rus'):
    """
    Функция, которая принимает строку text и возвращает вероятность того, насколько
    этот текст является реальным.
    МЕДЛЕННАЯ!
    """
    # Если в тексте нет слов
    if not text:
        return None

    text = text.lower()  # Преобразуем в нижний регистр
    words_list = re.findall(r'\w+', text)  # Разбиваем на слова


    total_words = len(words_list)
    real_words = 0

    # Подсчитываем количество реальных слов
    if lang == 'rus':
        for word in words_list:
            p: Parse = morph.parse(word)[0]
            if not (p.score < 0.5 or {'UNKN'} in p.tag
                    or isinstance(p.methods_stack[0][0], KnownSuffixAnalyzer.FakeDictionary)):
                real_words += 1
    else:
        for word in words_list:
            if wordnet.synsets(word):  # Очень медленная строчка!
                real_words += 1

    # Вычисляем вероятность как отношение реальных слов к общему количеству
    probability = real_words / total_words

    return probability


def read_txt(file_path):
    with open(file_path, 'r') as f:
        text = f.read()
    doc, index_map = create_document_and_map(text)
    return text, doc, index_map


def set_cell_border(cell: _Cell):
    """
    Устанавливает границы для ячейки таблицы (_Cell).
    Границы: черный цвет, толщина 12 (1.0pt), сплошная линия.
    Применяет стиль для всех сторон ячейки.
    """
    # Получаем XML представление ячейки
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Задаём стиль границ
    borders = {
        'top': {"sz": "2", "val": "single", "color": "000000", "space": "0"},
        'left': {"sz": "2", "val": "single", "color": "000000", "space": "0"},
        'bottom': {"sz": "2", "val": "single", "color": "000000", "space": "0"},
        'right': {"sz": "2", "val": "single", "color": "000000", "space": "0"}
    }

    # Применяем границы ко всем сторонам ячейки
    for border_name, border_attrs in borders.items():
        # Создаём или получаем существующий элемент для каждой стороны
        border_element = OxmlElement(f'w:{border_name}')
        for key, value in border_attrs.items():
            border_element.set(qn(f'w:{key}'), value)
        # Добавляем границу к ячейке
        tcPr.append(border_element)


def add_table_to_document(table: list[list[Any]],
                          doc: Document,
                          char_index: int,
                          full_text: list[str],
                          index_map: dict[int, Run]):
    full_text.append("\n")
    char_index += 1
    table_text = []
    doc_table = doc.add_table(rows=0, cols=len(table[0]))  # Создаём таблицу в DOCX
    for row_data in table:
        row_cells = doc_table.add_row().cells
        row_text = []
        for i, cell_text in enumerate(row_data):
            cell = row_cells[i]
            set_cell_border(cell)
            if cell_text:  # Если ячейка не пустая
                cell_paragraph = cell.paragraphs[0]
                for c in cell_text:
                    index_map[char_index] = cell_paragraph.add_run(str(c))
                    char_index += 1
                row_text.append(cell_text + ' ')
                char_index += 1
        table_text.append("".join(row_text) + '\n')
        char_index += 1

    full_text.append("".join(table_text))
    full_text.append("\n")
    char_index += 1
    return doc, char_index, full_text, index_map
