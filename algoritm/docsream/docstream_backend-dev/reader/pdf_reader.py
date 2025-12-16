from functools import lru_cache
from io import BytesIO

import fitz
import pdfplumber
import pymupdf
from PIL import Image
from docx import Document
from docx.text.run import Run
from pdfminer.pdfparser import PDFSyntaxError

from file_utils import get_image_text_rotate, create_document_and_map, add_table_to_document, get_image_text


@lru_cache(maxsize=128)
def read_pdf_with_tables(pdf_path) -> (str, Document, dict[int, Run]):
    full_text = []
    new_doc = Document()
    index_map: dict[int, Run] = dict()
    char_index = 0
    objects_on_pages: list[list[tuple[str: dict]]] = list()

    def add_table(table_data):
        nonlocal char_index, index_map, full_text, new_doc
        new_doc, char_index, full_text, index_map = add_table_to_document(table_data, new_doc, char_index, full_text, index_map)
    try:
        pdf = pdfplumber.open(pdf_path)
        # Открываем тот же PDF через PyMuPDF для работы с изображениями
        pdf_document = fitz.open(pdf_path)
        objects_on_page: list[tuple]
        for page in pdf.pages:
            tables_on_page = [(a,b) for a,b in zip(page.find_tables(), page.extract_tables())]
            text_on_page = page.extract_text_lines()
            objects_on_page = list()

            for table, table_data in tables_on_page:
                objects_on_page.append(('table', table, table_data))

            for line in text_on_page:
                skip = False
                if tables_on_page:
                    center = ((line['bottom'] + line['top']) / 2, (line['x0'] + line['x1']) / 2)
                    for table, _ in tables_on_page:
                        if table.bbox[1] <= center[0] <= table.bbox[3] \
                                and table.bbox[0] <= center[1] <= table.bbox[2]:
                            skip = True
                            break
                if skip:
                    continue
                objects_on_page.append(('text', line))

            for img in page.images:
                objects_on_page.append(('image', img))
                # objects_on_page.append(('image', {"bbox": bbox, "text": text}))

            def compare_key(obj):
                match obj[0]:
                    case 'text':
                        return (obj[1]['top'], obj[1]['x0'])
                    case 'table':
                        # bbox[0] - left
                        # bbox[1] - top
                        # bbox[2] - right
                        # bbox[3] - bottom
                        return (obj[1].bbox[1], obj[1].bbox[0])
                    case 'image':
                        return (obj[1]['y0'], obj[1]['x0'])

            objects_on_page.sort(key=compare_key)
            objects_on_pages.append(objects_on_page)

        pdf.close()

        for obj_list in objects_on_pages:
            for obj in obj_list:
                match obj[0]:
                    case 'text':
                        new_paragraph = new_doc.add_paragraph()
                        full_text.append(obj[1]['text'])
                        for c in obj[1]['text']:
                            index_map[char_index] = new_paragraph.add_run(str(c))
                            char_index += 1
                        full_text.append(' ')
                        char_index += 1
                    case 'table':
                        add_table(obj[2])
                    case 'image':
                        image_bytes = obj[1]['stream'].rawdata
                        text = get_image_text_rotate(image_bytes)
                        full_text.append(text)
                        new_paragraph = new_doc.add_paragraph()
                        for c in text:
                            index_map[char_index] = new_paragraph.add_run(str(c))
                            char_index += 1
        return ''.join(full_text), new_doc, index_map
    except PDFSyntaxError:
        return ''.join(full_text), new_doc, index_map



def read_pdf(file_path):
    doc = pymupdf.open(file_path)
    text = []
    for page_num in range(len(doc)):
        page = doc.load_page(page_num)
        text_part = page.get_text()
        text.append(text_part)

        image_list = page.get_images()
        for image in image_list:
            xref = image[0]
            img = doc.extract_image(xref)
            image_bytes = img["image"]
            text.append(get_image_text_rotate(image_bytes))

    text = ''.join(text)

    doc, index_map = create_document_and_map(text)
    return text, doc, index_map


def read_pdf_as_images(pdf_path):
    text = []
    # Открываем PDF
    pdf_document = fitz.open(pdf_path)

    # Проходимся по каждой странице
    for page_num in range(pdf_document.page_count):
        page = pdf_document.load_page(page_num)  # Загружаем страницу
        pix = page.get_pixmap(dpi=140)  # Получаем изображение страницы в виде объекта pixmap

        # Конвертируем pixmap в изображение PIL для дальнейшей работы
        image = Image.open(BytesIO(pix.tobytes(output="png")))

        # Сохраняем изображение в байты
        image_bytes_io = BytesIO()
        image.save(image_bytes_io, format="PNG")
        image_bytes = image_bytes_io.getvalue()

        # Передаем байты изображения в функцию для дальнейшей обработки
        text.append(get_image_text(image_bytes))

    text = ''.join(text)

    doc, index_map = create_document_and_map(text)
    return text, doc, index_map
