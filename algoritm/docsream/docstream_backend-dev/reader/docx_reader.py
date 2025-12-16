import aspose.words as aw
from docx import Document
from docx.oxml import CT_Tbl, CT_Tc
from docx.text.run import Run

from config import print
from file_utils import set_cell_border, add_table_to_document


def convert_table_to_text(block: CT_Tbl, doc: Document) -> (str, dict[int, Run]):
    char_index: int = 0

    # Словарь для сохранения каждого Run с его индексом
    index_map = {}

    # Список для полного текста
    full_text = []

    # Получаем количество строк и столбцов старой таблицы
    rows = block.tr_lst
    num_rows = len(rows)
    num_cols = len(rows[0].tc_lst) if num_rows > 0 else 0

    # Создаем новую таблицу в конце документа
    new_table = doc.add_table(rows=num_rows, cols=num_cols)

    # Проходим по строкам и столбцам старой таблицы
    for row_idx, row in enumerate(rows):
        cells = row.tc_lst
        cell: CT_Tc
        for col_idx, cell in enumerate(cells):
            # Получаем текст из старой ячейки (все абзацы)
            cell_text = '\n'.join([p.text for p in cell.p_lst]) + " "
            full_text.append(cell_text)  # Добавляем текст ячейки в список full_text

            # Создаем новый абзац в ячейке новой таблицы
            new_paragraph = new_table.cell(row_idx, col_idx).paragraphs[0]

            # Для каждого символа текста создаем отдельный Run
            for char in cell_text:
                run = new_paragraph.add_run(char)
                # Сохраняем Run в словарь с его индексом
                index_map[char_index] = run
                # Увеличиваем индекс для следующего символа
                char_index += 1
            set_cell_border(new_table.cell(row_idx, col_idx))

    # Соединяем весь текст ячеек в одну строку
    full_text_str = ''.join(full_text)

    return full_text_str, index_map


def read_docx(file_path, skip_first=0):
    """
    отличие от других функций только одно: сохранение стиля текста
    """
    doc = Document(file_path)
    new_doc = Document()
    full_text = []
    index_map: dict[int: Run] = dict()
    char_index = 0

    for paragraph in doc.paragraphs:
        if skip_first > 0:
            skip_first -= 1
            continue
        new_paragraph = new_doc.add_paragraph()
        new_paragraph.paragraph_format.alignment = paragraph.paragraph_format.alignment
        for run in paragraph.runs:
            for char in run.text:
                new_run = new_paragraph.add_run(char)
                new_run.font.highlight_color = run.font.highlight_color
                new_run.font.bold = run.font.bold
                new_run.font.italic = run.font.italic

                full_text.append(char)
                index_map[char_index] = new_run
                char_index += 1

    return ''.join(full_text), new_doc, index_map


def read_docx_with_tables(file_path, skip_first=0) -> (str, Document, dict[int, Run]):
    doc = Document(file_path)
    new_doc = Document()
    full_text = []
    index_map: dict[int: Run] = dict()
    char_index = 0

    def append_text(text, paragraph):
        nonlocal char_index, full_text
        for c in text:
            new_run = paragraph.add_run(c)
            index_map[char_index] = new_run
            char_index += 1
        full_text.append(text)

    def add_table(block: CT_Tbl):
        nonlocal char_index, full_text, new_doc, index_map
        rows = block.tr_lst
        table_data = []
        for row_idx, row in enumerate(rows):
            table_data.append([])
            cells = row.tc_lst
            cell: CT_Tc
            for col_idx, cell in enumerate(cells):
                cell_text = '\n'.join([p.text for p in cell.p_lst])
                table_data[-1].append(cell_text)

        new_doc, char_index, full_text, index_map = add_table_to_document(table_data, new_doc, char_index, full_text, index_map)

    block: CT_Tbl
    for block in doc.element.body:
        if skip_first > 0:
            skip_first -= 1
            continue

        if isinstance(block, str):
            # block_type = "Текст"
            new_paragraph = new_doc.add_paragraph()
            append_text(block, new_paragraph)
        elif block.tag.endswith('p'):
            # block_type = "Абзац"
            new_paragraph = new_doc.add_paragraph()
            append_text(block.text, new_paragraph)
        elif block.tag.endswith('tbl'):
            # block_type = "Таблица"
            add_table(block)
        elif block.tag.endswith('drawing'):
            # block_type = "Картинка"
            # new_doc.add_picture() # skipped
            print(f"Пропущено изображение: {block.tag}")
        elif block.tag.endswith('sectPr'):
            pass
        else:
            print("Неизвестный элемент: " + block.tag)

    return ''.join(full_text), new_doc, index_map


def read_doc(file_path):
    doc = aw.Document(file_path)
    file_path += 'x'
    doc.save(file_path)
    return read_docx(file_path, skip_first=1)
